import io
import os
import json


from datetime import datetime
from zoneinfo import ZoneInfo
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_file
from pymongo import MongoClient
from bson.objectid import ObjectId
from werkzeug.security import generate_password_hash, check_password_hash
from openai import OpenAI
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase.pdfmetrics import stringWidth
from dotenv import load_dotenv
try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Inches, Pt
except ImportError:
    Document = None

load_dotenv(override=True)

DISPLAY_TIMEZONE = ZoneInfo("America/Los_Angeles")


def get_required_env(name):
    value = os.environ.get(name)
    if not value:
        raise RuntimeError(f"{name} is not set.")
    return value


def normalize_text_block(value):
    if value is None:
        return ""
    return str(value).strip()


def build_export_filename(profile_name, extension):
    raw_name = normalize_text_block(profile_name) or "targetresume"
    safe_name = "_".join(raw_name.split())
    safe_name = "".join(ch for ch in safe_name if ch.isalnum() or ch in ("_", "-")).strip("_-")
    if not safe_name:
        safe_name = "targetresume"
    return f"{safe_name}_targetresume.{extension}"


def parse_bullets(raw_value):
    if raw_value is None:
        return []
    if isinstance(raw_value, list):
        values = raw_value
    else:
        values = str(raw_value).splitlines()

    bullets = []
    for value in values:
        cleaned = str(value).strip()
        if not cleaned:
            continue
        if cleaned.startswith(("- ", "* ")):
            cleaned = cleaned[2:].strip()
        bullets.append(cleaned)
    return bullets


def normalize_resume_entries(raw_value):
    if not raw_value:
        return []

    if isinstance(raw_value, str):
        try:
            raw_value = json.loads(raw_value)
        except json.JSONDecodeError:
            return []

    if not isinstance(raw_value, list):
        return []

    entries = []
    for item in raw_value:
        if not isinstance(item, dict):
            continue

        title = normalize_text_block(item.get("title"))
        details = normalize_text_block(item.get("details"))
        location = normalize_text_block(item.get("location"))
        dates = normalize_text_block(item.get("dates"))
        bullets = parse_bullets(item.get("bullets"))

        if title or details or location or dates or bullets:
            entries.append({
                "title": title,
                "details": details,
                "location": location,
                "dates": dates,
                "bullets": bullets
            })

    return entries


def format_resume_entries(entries):
    sections = []
    for entry in entries:
        title = normalize_text_block(entry.get("title"))
        details = normalize_text_block(entry.get("details"))
        location = normalize_text_block(entry.get("location"))
        dates = normalize_text_block(entry.get("dates"))
        bullets = parse_bullets(entry.get("bullets"))

        header_parts = [part for part in [title, details or location, dates] if part]
        lines = []
        if header_parts:
            lines.append(" | ".join(header_parts))
        lines.extend(f"- {bullet}" for bullet in bullets)

        if lines:
            sections.append("\n".join(lines))

    return "\n\n".join(sections)


def normalize_skills_entries(raw_value):
    if not raw_value:
        return []

    if isinstance(raw_value, str):
        try:
            raw_value = json.loads(raw_value)
        except json.JSONDecodeError:
            return []

    if not isinstance(raw_value, list):
        return []

    entries = []
    for item in raw_value:
        if not isinstance(item, dict):
            continue

        category = normalize_text_block(item.get("category"))
        values = item.get("values", [])
        if isinstance(values, str):
            values = [part.strip() for part in values.split(",")]

        normalized_values = [normalize_text_block(value) for value in values if normalize_text_block(value)]
        if category or normalized_values:
            entries.append({
                "category": category,
                "values": normalized_values
            })

    return entries


def format_skills_entries(entries):
    lines = []
    for entry in entries:
        category = normalize_text_block(entry.get("category"))
        values = [normalize_text_block(value) for value in entry.get("values", []) if normalize_text_block(value)]
        if not category and not values:
            continue
        if category and values:
            lines.append(f"- {category}: {', '.join(values)}")
        elif category:
            lines.append(f"- {category}")
        else:
            lines.append(f"- {', '.join(values)}")
    return "\n".join(lines)


def normalize_certification_entries(raw_value):
    if not raw_value:
        return []

    if isinstance(raw_value, str):
        try:
            raw_value = json.loads(raw_value)
        except json.JSONDecodeError:
            return []

    if not isinstance(raw_value, list):
        return []

    entries = []
    for item in raw_value:
        if not isinstance(item, dict):
            continue

        name = normalize_text_block(item.get("name"))
        date = normalize_text_block(item.get("date"))
        description = normalize_text_block(item.get("description"))

        if name or date or description:
            entries.append({
                "name": name,
                "date": date,
                "description": description
            })

    return entries


def format_certification_entries(entries):
    sections = []
    for entry in entries:
        name = normalize_text_block(entry.get("name"))
        date = normalize_text_block(entry.get("date"))
        description = normalize_text_block(entry.get("description"))
        header_parts = [part for part in [name, date] if part]
        lines = []
        if header_parts:
            lines.append(" | ".join(header_parts))
        if description:
            lines.append(description)
        if lines:
            sections.append("\n".join(lines))
    return "\n\n".join(sections)


def parse_certifications_text_to_entries(text):
    entries = []
    current = None
    for raw_line in normalize_text_block(text).splitlines():
        line = raw_line.strip()
        if not line:
            if current:
                entries.append(current)
                current = None
            continue

        if current is None:
            parts = [part.strip() for part in line.split("|")]
            current = {
                "name": parts[0] if parts else "",
                "date": " | ".join(parts[1:]) if len(parts) > 1 else "",
                "description": ""
            }
        else:
            current["description"] = f"{current['description']} {line}".strip() if current.get("description") else line

    if current:
        entries.append(current)

    return [entry for entry in entries if entry.get("name") or entry.get("date") or entry.get("description")]


def parse_skills_text_to_entries(text):
    entries = []
    for raw_line in normalize_text_block(text).splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("- "):
            line = line[2:].strip()
        parts = line.split(":")
        if len(parts) > 1:
            category = parts[0].strip()
            values = [item.strip() for item in ":".join(parts[1:]).split(",") if item.strip()]
        else:
            category = ""
            values = [item.strip() for item in line.split(",") if item.strip()]
        if category or values:
            entries.append({
                "category": category,
                "values": values
            })
    return entries


def parse_resume_text_to_entries(text):
    entries = []
    current = None
    for raw_line in normalize_text_block(text).splitlines():
        line = raw_line.strip()
        if not line:
            if current:
                entries.append(current)
                current = None
            continue
        if line.startswith("- "):
            if current is None:
                current = {"title": "", "details": "", "bullets": []}
            current["bullets"].append(line[2:].strip())
        else:
            if current:
                entries.append(current)
            parts = [part.strip() for part in line.split("|")]
            details = " | ".join(parts[1:]) if len(parts) > 1 else ""
            location = ""
            dates = ""
            if len(parts) >= 3:
                location = parts[1]
                dates = " | ".join(parts[2:])
                details = ""
            current = {
                "title": parts[0] if parts else "",
                "details": details,
                "location": location,
                "dates": dates,
                "bullets": []
            }
    if current:
        entries.append(current)
    return [entry for entry in entries if entry.get("title") or entry.get("details") or entry.get("location") or entry.get("dates") or entry.get("bullets")]


def prepare_profile_for_view(profile):
    profile = profile or {}
    profile["skills_entries"] = normalize_skills_entries(profile.get("skills_entries")) or parse_skills_text_to_entries(profile.get("skills", ""))
    profile["projects_entries"] = normalize_resume_entries(profile.get("projects_entries")) or parse_resume_text_to_entries(profile.get("projects", ""))
    profile["experience_entries"] = normalize_resume_entries(profile.get("experience_entries")) or parse_resume_text_to_entries(profile.get("experience", ""))
    profile["certifications_entries"] = normalize_certification_entries(profile.get("certifications_entries")) or parse_certifications_text_to_entries(profile.get("certifications", ""))
    return profile


def prepare_resume_for_view(resume):
    resume = resume or {}
    resume["skills_entries"] = normalize_skills_entries(resume.get("skills_entries")) or parse_skills_text_to_entries(resume.get("skills", ""))
    resume["projects_entries"] = normalize_resume_entries(resume.get("projects_entries")) or parse_resume_text_to_entries(resume.get("projects", ""))
    resume["experience_entries"] = normalize_resume_entries(resume.get("experience_entries")) or parse_resume_text_to_entries(resume.get("experience", ""))
    resume["certifications_entries"] = normalize_certification_entries(resume.get("certifications_entries")) or parse_certifications_text_to_entries(resume.get("certifications", ""))
    return resume


def format_datetime_for_display(value):
    if not value:
        return "", ""

    display_value = value
    if value.tzinfo is None:
        display_value = value.replace(tzinfo=ZoneInfo("UTC"))

    display_value = display_value.astimezone(DISPLAY_TIMEZONE)
    return display_value.strftime("%b %d, %Y"), display_value.strftime("%I:%M %p")


def ensure_folder_record(user_id, folder_name):
    folder_name = (folder_name or "").strip()
    if not folder_name or folder_name == "Saved Drafts":
        return

    highest_order_folder = folders_collection.find_one(
        {"user_id": user_id},
        sort=[("sort_order", -1)]
    )
    next_order = (highest_order_folder.get("sort_order", 0) + 1) if highest_order_folder else 1

    folders_collection.update_one(
        {"user_id": user_id, "name": folder_name},
        {"$setOnInsert": {
            "user_id": user_id,
            "name": folder_name,
            "sort_order": next_order,
            "created_at": datetime.utcnow()
        }},
        upsert=True
    )


def build_folder_sidebar(user_id, resumes):
    folder_docs = list(
        folders_collection.find({"user_id": user_id}).sort("sort_order", 1)
    )
    ordered_folder_names = []
    seen = set()

    for folder in folder_docs:
        name = folder.get("name", "").strip()
        if not name or name in seen:
            continue
        ordered_folder_names.append(name)
        seen.add(name)

    resume_folder_names = {
        (resume.get("folder") or "Saved Drafts").strip() or "Saved Drafts"
        for resume in resumes
    }

    for folder_name in sorted(resume_folder_names):
        if folder_name not in seen and folder_name != "Saved Drafts":
            ordered_folder_names.append(folder_name)
            seen.add(folder_name)

    folder_items = [{"name": "Saved Drafts", "can_manage": False}]
    folder_items.extend({
        "name": name,
        "can_manage": True
    } for name in ordered_folder_names if name != "Saved Drafts")
    return folder_items


def parse_ai_rewrite_response(parsed, profile):
    skills_entries = normalize_skills_entries(parsed.get("skills_entries"))
    project_entries = normalize_resume_entries(parsed.get("projects_entries"))
    experience_entries = normalize_resume_entries(parsed.get("experience_entries"))
    certifications_entries = normalize_certification_entries(parsed.get("certifications_entries"))
    fallback_projects = normalize_resume_entries(profile.get("projects_entries"))

    if not skills_entries:
        skills_entries = normalize_skills_entries(profile.get("skills_entries"))
    if not project_entries:
        project_entries = fallback_projects
    if not experience_entries:
        experience_entries = normalize_resume_entries(profile.get("experience_entries"))
    if not certifications_entries:
        certifications_entries = normalize_certification_entries(profile.get("certifications_entries"))

    if len(project_entries) < 3 and fallback_projects:
        seen = {
            (
                normalize_text_block(entry.get("title")),
                normalize_text_block(entry.get("details"))
            )
            for entry in project_entries
        }
        for fallback_entry in fallback_projects:
            key = (
                normalize_text_block(fallback_entry.get("title")),
                normalize_text_block(fallback_entry.get("details"))
            )
            if key in seen:
                continue
            project_entries.append(fallback_entry)
            seen.add(key)
            if len(project_entries) >= 3:
                break

    project_entries = [
        {
            **entry,
            "bullets": parse_bullets(entry.get("bullets"))[:4]
        }
        for entry in project_entries[:4]
    ]
    experience_entries = [
        {
            **entry,
            "bullets": parse_bullets(entry.get("bullets"))[:4]
        }
        for entry in experience_entries
    ]

    return {
        "skills": format_skills_entries(skills_entries) or profile.get("skills", ""),
        "projects": format_resume_entries(project_entries) or profile.get("projects", ""),
        "experience": format_resume_entries(experience_entries) or profile.get("experience", ""),
        "certifications": format_certification_entries(certifications_entries) or profile.get("certifications", ""),
        "skills_entries": skills_entries,
        "projects_entries": project_entries,
        "experience_entries": experience_entries,
        "certifications_entries": certifications_entries
    }


def split_text_to_lines(text, font_name, font_size, max_width):
    text = normalize_text_block(text)
    if not text:
        return []

    lines = []
    for paragraph in text.splitlines():
        paragraph = paragraph.strip()
        if not paragraph:
            lines.append("")
            continue

        words = paragraph.split()
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            if stringWidth(candidate, font_name, font_size) <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)

    return lines


def looks_like_date_text(text):
    text = normalize_text_block(text).lower()
    if not text:
        return False

    month_tokens = [
        "jan", "feb", "mar", "apr", "may", "jun",
        "jul", "aug", "sep", "oct", "nov", "dec",
        "present", "current"
    ]
    return any(token in text for token in month_tokens) or any(char.isdigit() for char in text)


app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "targetresume_dev_secret_key")

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
openai_client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
GA_MEASUREMENT_ID = os.environ.get("GA_MEASUREMENT_ID", "").strip()

MONGODB_URI = get_required_env("MONGODB_URI")

client = MongoClient(MONGODB_URI)

db = client["TargetResume"]
users_collection = db["users"]
resumes_collection = db["resumes"]
jobs_collection = db["job_tracker"]
profiles_collection = db["profiles"]
folders_collection = db["folders"]

HOST = os.environ.get("HOST", "0.0.0.0")
PORT = int(os.environ.get("PORT", "5025"))


@app.context_processor
def inject_global_template_vars():
    return {
        "ga_measurement_id": GA_MEASUREMENT_ID
    }


@app.route("/")
def home():
    if "user_id" in session:
        return redirect(url_for("dashboard"))
    return redirect(url_for("login"))


@app.route("/dashboard")
def dashboard():
    if "user_id" not in session:
        return redirect(url_for("login"))

    user_id = session["user_id"]
    profile = prepare_profile_for_view(profiles_collection.find_one({"user_id": user_id}))
    resume_id = request.args.get("resume_id", "").strip()
    selected_resume = None

    if resume_id:
        selected_resume = resumes_collection.find_one({
            "_id": ObjectId(resume_id),
            "user_id": user_id
        })
        selected_resume = prepare_resume_for_view(selected_resume)

    return render_template("dashboard.html", profile=profile, selected_resume=selected_resume)


@app.route("/generate-resume-preview", methods=["POST"])
def generate_resume_preview():
    if "user_id" not in session:
        return jsonify({"error": "Unauthorized"}), 401

    user_id = session["user_id"]
    profile = profiles_collection.find_one({"user_id": user_id}) or {}
    skills_entries = normalize_skills_entries(profile.get("skills_entries"))
    project_entries = normalize_resume_entries(profile.get("projects_entries"))
    experience_entries = normalize_resume_entries(profile.get("experience_entries"))
    certifications_entries = normalize_certification_entries(profile.get("certifications_entries"))

    education_school = ""
    if profile.get("school"):
        education_school += profile.get("school", "")
    if profile.get("school_location"):
        education_school += f", {profile.get('school_location')}"

    response = {
        "name": profile.get("name", "Your Name"),
        "contact": " | ".join(filter(None, [
            profile.get("email"),
            profile.get("phone"),
            profile.get("linkedin"),
            profile.get("portfolio")
        ])) or "Email | Phone | LinkedIn | Portfolio",
        "education_top": education_school,
        "education_grad": profile.get("expected_grad", ""),
        "education_bottom": profile.get("degree", ""),
        "skills": format_skills_entries(skills_entries) or profile.get("skills", ""),
        "projects": format_resume_entries(project_entries) or profile.get("projects", ""),
        "experience": format_resume_entries(experience_entries) or profile.get("experience", ""),
        "certifications": format_certification_entries(certifications_entries) or profile.get("certifications", ""),
        "skills_entries": skills_entries,
        "projects_entries": project_entries,
        "experience_entries": experience_entries,
        "certifications_entries": certifications_entries
    }

    return jsonify(response)

@app.route("/ai-rewrite-preview", methods=["POST"])
def ai_rewrite_preview():
    if "user_id" not in session:
        return jsonify({"error": "Unauthorized"}), 401

    if not openai_client:
        return jsonify({"error": "OPENAI_API_KEY is not set."}), 500

    user_id = session["user_id"]
    profile = prepare_profile_for_view(profiles_collection.find_one({"user_id": user_id}))

    job_title = request.form.get("job_title", "").strip()
    job_description = request.form.get("job_description", "").strip()
    notes = request.form.get("notes", "").strip()

    if not job_description:
        return jsonify({"error": "Job description is required."}), 400

    education_text = ""
    if profile.get("school"):
        education_text += profile.get("school", "")
    if profile.get("school_location"):
        education_text += f", {profile.get('school_location')}"
    if profile.get("expected_grad"):
        education_text += f" Graduation {profile.get('expected_grad')}"
    if profile.get("degree"):
        education_text += f"\n{profile.get('degree')}"

    source_skills = profile.get("skills_entries") or profile.get("skills", "")
    source_projects = profile.get("projects_entries") or profile.get("projects", "")
    source_experience = profile.get("experience_entries") or profile.get("experience", "")

    prompt = f"""
You are helping tailor a resume for a specific job.

Use ONLY the user's provided information.
Do NOT invent employers, projects, dates, degrees, metrics, or skills.
Do NOT add tools, technologies, certifications, coursework, or achievements that are not explicitly present in the user's profile.
Do NOT rewrite facts into stronger-sounding facts if that changes meaning.
Keep organization names, role names, dates, and locations accurate to the user's input.
Do NOT create a professional summary.
Rewrite only these sections:
skills
certifications
projects
experience

The output must fit a strong one-page student resume.
Select only the strongest and most relevant details for this specific target job.
Do not include everything if that would make the resume too long.
Prefer concise, high-impact bullets.
Prefer stronger existing bullets over rewriting them into vague filler.
Avoid repeating the same accomplishment across multiple projects or jobs.
If a section is weak for this target job, return fewer stronger entries rather than forcing extra filler.
Return at least 3 projects when the user has 3 or more available.
Order projects from strongest and most relevant first to weakest and least relevant last.
Keep 2-4 bullets per project or experience entry.
If an entry has 4 strong relevant bullets, keep all 4 instead of shortening it unnecessarily.
Only include certifications that are genuinely relevant or meaningfully strengthen the application.
If the user has fewer than 3 projects total, return only the real projects they actually have.

IMPORTANT:
Return valid JSON.

Return exactly this JSON format:
{{
  "skills_entries": [
    {{
      "category": "Languages",
      "values": ["Java", "Python", "JavaScript"]
    }}
  ],
  "projects_entries": [
    {{
      "title": "Project name",
      "details": "optional short detail line such as tech stack",
      "bullets": [
        "short bullet",
        "short bullet"
      ]
    }}
  ],
  "certifications_entries": [
    {{
      "name": "Certification name",
      "date": "optional date obtained",
      "description": "optional short description"
    }}
  ],
  "experience_entries": [
    {{
      "title": "Role or organization",
      "details": "optional short detail line such as company | location | dates",
      "bullets": [
        "short bullet",
        "short bullet"
      ]
    }}
  ]
}}

Do not return any keys other than:
skills_entries
certifications_entries
projects_entries
experience_entries

Target job title:
{job_title}

Target job description:
{job_description}

Special focus from user:
{notes}

User profile data:
Name: {profile.get("name", "")}
Education: {education_text}
Skills: {json.dumps(source_skills, ensure_ascii=True)}
Projects: {json.dumps(source_projects, ensure_ascii=True)}
Experience: {json.dumps(source_experience, ensure_ascii=True)}
Certifications: {json.dumps(profile.get("certifications_entries") or profile.get("certifications", ""), ensure_ascii=True)}
"""

    try:
        response = openai_client.responses.create(
            model="gpt-5.4-mini",
            input=prompt
        )

        raw_text = response.output_text.strip()
        parsed = json.loads(raw_text)
        return jsonify(parse_ai_rewrite_response(parsed, profile))

    except Exception as e:
        return jsonify({"error": str(e)}), 500


def build_resume_document(user_id, profile, form_data):
    job_title = form_data.get("job_title", "").strip()
    folder = form_data.get("folder", "").strip() or "Saved Drafts"

    return {
        "user_id": user_id,
        "title": job_title if job_title else "Untitled Resume",
        "folder": folder,
        "job_title": job_title,
        "job_description": form_data.get("job_description"),
        "notes": form_data.get("notes"),
        "name": profile.get("name", ""),
        "email": profile.get("email", ""),
        "phone": profile.get("phone", ""),
        "linkedin": profile.get("linkedin", ""),
        "portfolio": profile.get("portfolio", ""),
        "school": profile.get("school", ""),
        "school_location": profile.get("school_location", ""),
        "expected_grad": profile.get("expected_grad", ""),
        "degree": profile.get("degree", ""),
        "skills": form_data.get("tailored_skills"),
        "projects": form_data.get("tailored_projects"),
        "experience": form_data.get("tailored_experience"),
        "certifications": form_data.get("tailored_certifications"),
        "skills_entries": normalize_skills_entries(form_data.get("skills_entries")),
        "projects_entries": normalize_resume_entries(form_data.get("projects_entries")),
        "experience_entries": normalize_resume_entries(form_data.get("experience_entries")),
        "certifications_entries": normalize_certification_entries(form_data.get("certifications_entries"))
    }


@app.route("/save-resume-version", methods=["POST"])
def save_resume_version():
    if "user_id" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401

    user_id = session["user_id"]
    profile = profiles_collection.find_one({"user_id": user_id}) or {}
    resume_doc = build_resume_document(user_id, profile, request.form)
    resume_doc["created_at"] = datetime.utcnow()
    resume_doc["updated_at"] = datetime.utcnow()

    ensure_folder_record(user_id, resume_doc["folder"])

    result = resumes_collection.insert_one(resume_doc)

    return jsonify({
        "success": True,
        "resume_id": str(result.inserted_id)
    })


@app.route("/update-resume-version/<resume_id>", methods=["POST"])
def update_resume_version(resume_id):
    if "user_id" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401

    user_id = session["user_id"]
    existing_resume = resumes_collection.find_one({
        "_id": ObjectId(resume_id),
        "user_id": user_id
    })

    if not existing_resume:
        return jsonify({"success": False, "error": "Resume not found."}), 404

    profile = profiles_collection.find_one({"user_id": user_id}) or {}
    updated_resume = build_resume_document(user_id, profile, request.form)
    updated_resume["title"] = existing_resume.get("title") or updated_resume["title"]
    updated_resume["updated_at"] = datetime.utcnow()

    ensure_folder_record(user_id, updated_resume["folder"])

    resumes_collection.update_one(
        {"_id": existing_resume["_id"], "user_id": user_id},
        {"$set": updated_resume}
    )

    return jsonify({
        "success": True,
        "resume_id": str(existing_resume["_id"])
    })


@app.route("/export-resume")
def export_resume():
    if "user_id" not in session:
        return redirect(url_for("login"))

    user_id = session["user_id"]
    profile = profiles_collection.find_one({"user_id": user_id}) or {}
    skills_entries = normalize_skills_entries(profile.get("skills_entries"))
    project_entries = normalize_resume_entries(profile.get("projects_entries"))
    experience_entries = normalize_resume_entries(profile.get("experience_entries"))
    certifications_entries = normalize_certification_entries(profile.get("certifications_entries"))

    tailored_skills = request.args.get(
        "tailored_skills",
        format_skills_entries(skills_entries) or profile.get("skills", "")
    )
    tailored_projects = request.args.get(
        "tailored_projects",
        format_resume_entries(project_entries) or profile.get("projects", "")
    )
    tailored_experience = request.args.get(
        "tailored_experience",
        format_resume_entries(experience_entries) or profile.get("experience", "")
    )
    certifications_text = request.args.get(
        "tailored_certifications",
        format_certification_entries(certifications_entries) or profile.get("certifications", "")
    )

    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    page_width, page_height = letter
    left_margin = 36
    right_margin = 36
    content_width = page_width - left_margin - right_margin
    y = page_height - 36
    bottom_margin = 36

    education_school = ""
    if profile.get("school"):
        education_school += profile.get("school", "")
    if profile.get("school_location"):
        education_school += f", {profile.get('school_location')}"
    education_grad = profile.get("expected_grad", "")
    education_degree = profile.get("degree", "")

    skills_lines = split_text_to_lines(tailored_skills.replace("\u2022", "-"), "Times-Roman", 10.5, content_width - 14)

    def new_page():
        nonlocal y
        p.showPage()
        y = page_height - 42

    def ensure_space(required_height):
        nonlocal y
        if y - required_height < bottom_margin:
            new_page()

    def draw_section_title(title):
        nonlocal y
        ensure_space(30)
        y -= 4
        p.setFont("Times-Bold", 11)
        p.drawString(left_margin, y, title)
        y -= 6
        p.setLineWidth(0.8)
        p.line(left_margin, y, page_width - right_margin, y)
        y -= 14

    def draw_wrapped_line(text, font_name="Times-Roman", font_size=10.5, indent=0, leading=13):
        nonlocal y
        max_width = content_width - indent
        lines = split_text_to_lines(text, font_name, font_size, max_width)
        for line in lines:
            ensure_space(leading)
            p.setFont(font_name, font_size)
            p.drawString(left_margin + indent, y, line)
            y -= leading

    def draw_bullet_list(lines):
        nonlocal y
        bullet_indent = 10
        text_indent = 20
        for bullet in lines:
            clean_bullet = bullet.strip()
            if not clean_bullet:
                continue
            if clean_bullet.startswith("- "):
                clean_bullet = clean_bullet[2:].strip()

            wrapped = split_text_to_lines(clean_bullet, "Times-Roman", 10.5, content_width - text_indent)
            if not wrapped:
                continue

            ensure_space(13 * len(wrapped))
            p.setFont("Times-Roman", 10.5)
            p.drawString(left_margin + bullet_indent, y, u"\u2022")
            p.drawString(left_margin + text_indent, y, wrapped[0])
            y -= 13
            for continuation in wrapped[1:]:
                ensure_space(13)
                p.drawString(left_margin + text_indent, y, continuation)
                y -= 13

    def draw_two_column_skills(lines):
        nonlocal y
        cleaned_lines = []
        for line in lines:
            clean = line.strip()
            if not clean:
                continue
            if clean.startswith("- "):
                clean = clean[2:].strip()
            category = ""
            values = clean
            if ":" in clean:
                category, values = clean.split(":", 1)
                category = category.strip()
                values = values.strip()
            cleaned_lines.append({
                "category": category,
                "values": values
            })

        if not cleaned_lines:
            return

        column_gap = 22
        column_width = (content_width - column_gap) / 2
        left_lines = cleaned_lines[::2]
        right_lines = cleaned_lines[1::2]
        row_count = max(len(left_lines), len(right_lines))

        for row_index in range(row_count):
            left_entry = left_lines[row_index] if row_index < len(left_lines) else None
            right_entry = right_lines[row_index] if row_index < len(right_lines) else None

            left_text = f"{left_entry['category']}: {left_entry['values']}" if left_entry and left_entry["category"] else (left_entry["values"] if left_entry else "")
            right_text = f"{right_entry['category']}: {right_entry['values']}" if right_entry and right_entry["category"] else (right_entry["values"] if right_entry else "")

            left_wrapped = split_text_to_lines(left_text, "Times-Roman", 10.2, column_width - 14) if left_text else []
            right_wrapped = split_text_to_lines(right_text, "Times-Roman", 10.2, column_width - 14) if right_text else []
            row_height = max(len(left_wrapped), len(right_wrapped), 1) * 12
            ensure_space(row_height)

            def draw_skill_column(entry, x_start, wrapped_lines):
                if not entry or not wrapped_lines:
                    return

                for index, line in enumerate(wrapped_lines):
                    current_y = y - (index * 12)
                    if index == 0:
                        p.setFont("Times-Roman", 10.2)
                        p.drawString(x_start, current_y, u"\u2022")

                        text_x = x_start + 10
                        if entry["category"]:
                            category_text = f"{entry['category']}:"
                            p.setFont("Times-Bold", 10.2)
                            p.drawString(text_x, current_y, category_text)
                            category_width = stringWidth(category_text, "Times-Bold", 10.2)
                            values_text = f" {entry['values']}" if entry["values"] else ""
                            if values_text:
                                p.setFont("Times-Roman", 10.2)
                                p.drawString(text_x + category_width, current_y, values_text)
                        else:
                            p.setFont("Times-Roman", 10.2)
                            p.drawString(text_x, current_y, line)
                    else:
                        p.setFont("Times-Roman", 10.2)
                        p.drawString(x_start + 10, current_y, line)

            draw_skill_column(left_entry, left_margin, left_wrapped)

            right_x = left_margin + column_width + column_gap
            draw_skill_column(right_entry, right_x, right_wrapped)

            y -= row_height

    def parse_structured_text(text):
        sections = []
        current = None
        for raw_line in text.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith("- "):
                if current is None:
                    current = {"header": "", "bullets": []}
                current["bullets"].append(line)
            else:
                if current:
                    sections.append(current)
                current = {"header": line, "bullets": []}
        if current:
            sections.append(current)
        return sections

    def draw_entry_sections(section_text):
        nonlocal y
        for entry in parse_structured_text(section_text):
            header = entry["header"]
            bullets = entry["bullets"]
            if header:
                header_parts = [part.strip() for part in header.split("|")]
                title = header_parts[0] if header_parts else ""
                details = " | ".join(header_parts[1:]) if len(header_parts) > 1 else ""
                draw_wrapped_line(title, font_name="Times-Bold", font_size=10.5, leading=12)
                if details:
                    draw_wrapped_line(details, font_name="Times-Italic", font_size=10, leading=12)
            draw_bullet_list(bullets)
            y -= 3

    def draw_experience_sections(section_text):
        nonlocal y
        for entry in parse_structured_text(section_text):
            header = entry["header"]
            bullets = entry["bullets"]

            if header:
                header_parts = [part.strip() for part in header.split("|") if part.strip()]
                left_text = ""
                right_text = ""

                if header_parts:
                    if len(header_parts) == 1:
                        left_text = header_parts[0]
                    elif len(header_parts) == 2 and not looks_like_date_text(header_parts[1]):
                        left_text = ", ".join(header_parts)
                    else:
                        left_text = ", ".join(header_parts[:-1])
                        right_text = header_parts[-1]

                left_font_name = "Times-Bold"
                left_font_size = 9.8
                right_font_name = "Times-Italic"
                right_font_size = 9.8
                right_width = stringWidth(right_text, right_font_name, right_font_size) if right_text else 0
                left_width = content_width - right_width - 18 if right_text else content_width
                left_lines = split_text_to_lines(left_text, left_font_name, left_font_size, left_width) if left_text else []
                line_count = max(len(left_lines), 1)
                ensure_space(line_count * 12)

                p.setFont(left_font_name, left_font_size)
                if left_lines:
                    p.drawString(left_margin, y, left_lines[0])
                if right_text:
                    p.setFont(right_font_name, right_font_size)
                    p.drawRightString(page_width - right_margin, y, right_text)
                y -= 12

                for continuation in left_lines[1:]:
                    ensure_space(12)
                    p.setFont(left_font_name, left_font_size)
                    p.drawString(left_margin, y, continuation)
                    y -= 12

            draw_bullet_list(bullets)
            y -= 3

    def draw_certification_sections(entries):
        nonlocal y
        for entry in entries:
            name = normalize_text_block(entry.get("name"))
            date = normalize_text_block(entry.get("date"))
            description = normalize_text_block(entry.get("description"))
            if name or date:
                ensure_space(12)
                p.setFont("Times-Roman", 10.5)
                if name:
                    p.drawString(left_margin, y, u"\u2022")
                p.setFont("Times-Bold", 10.5)
                if name:
                    p.drawString(left_margin + 10, y, name)
                if date:
                    p.setFont("Times-Italic", 10)
                    p.drawRightString(page_width - right_margin, y, date)
                y -= 12
            if description:
                draw_wrapped_line(description, font_name="Times-Roman", font_size=10.3, indent=10, leading=12)
            y -= 3

    p.setFont("Times-Bold", 18)
    p.drawCentredString(page_width / 2, y, profile.get("name", "Your Name"))
    y -= 20

    contact = " | ".join(filter(None, [
        profile.get("email"),
        profile.get("phone"),
        profile.get("linkedin"),
        profile.get("portfolio")
    ]))
    if contact:
        for line in split_text_to_lines(contact, "Times-Roman", 10.5, content_width):
            p.setFont("Times-Roman", 10.5)
            p.drawCentredString(page_width / 2, y, line)
            y -= 13
    y -= 8

    if education_school or education_grad or education_degree:
        draw_section_title("EDUCATION")
        ensure_space(28)
        p.setFont("Times-Italic", 10.5)
        if education_school:
            p.drawString(left_margin, y, education_school)
        if education_grad:
            grad_text = f"Graduation {education_grad}"
            p.drawRightString(page_width - right_margin, y, grad_text)
        y -= 13
        if education_degree:
            p.setFont("Times-Roman", 10.5)
            p.drawString(left_margin, y, u"\u2022")
            p.drawString(left_margin + 10, y, education_degree)
            y -= 16

    if skills_lines:
        draw_section_title("TECHNICAL SKILLS")
        draw_two_column_skills(skills_lines)
        y -= 4

    if certifications_entries or certifications_text:
        draw_section_title("CERTIFICATIONS")
        draw_certification_sections(
            certifications_entries if certifications_entries else parse_certifications_text_to_entries(certifications_text)
        )

    if normalize_text_block(tailored_projects):
        draw_section_title("PROJECTS")
        draw_entry_sections(tailored_projects)

    if normalize_text_block(tailored_experience):
        draw_section_title("WORK EXPERIENCE")
        draw_experience_sections(tailored_experience)

    p.save()
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name=build_export_filename(profile.get("name", ""), "pdf"),
        mimetype="application/pdf"
    )


@app.route("/export-resume-docx")
def export_resume_docx():
    if "user_id" not in session:
        return redirect(url_for("login"))

    if Document is None:
        return "DOCX export is unavailable because python-docx is not installed.", 500

    user_id = session["user_id"]
    profile = profiles_collection.find_one({"user_id": user_id}) or {}
    skills_entries = normalize_skills_entries(profile.get("skills_entries"))
    project_entries = normalize_resume_entries(profile.get("projects_entries"))
    experience_entries = normalize_resume_entries(profile.get("experience_entries"))
    certifications_entries = normalize_certification_entries(profile.get("certifications_entries"))

    tailored_skills = request.args.get(
        "tailored_skills",
        format_skills_entries(skills_entries) or profile.get("skills", "")
    )
    tailored_projects = request.args.get(
        "tailored_projects",
        format_resume_entries(project_entries) or profile.get("projects", "")
    )
    tailored_experience = request.args.get(
        "tailored_experience",
        format_resume_entries(experience_entries) or profile.get("experience", "")
    )
    certifications_text = request.args.get(
        "tailored_certifications",
        format_certification_entries(certifications_entries) or profile.get("certifications", "")
    )

    education_school = ""
    if profile.get("school"):
        education_school += profile.get("school", "")
    if profile.get("school_location"):
        education_school += f", {profile.get('school_location')}"
    education_grad = profile.get("expected_grad", "")
    education_degree = profile.get("degree", "")

    def parse_structured_text(text):
        sections = []
        current = None
        for raw_line in text.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith("- "):
                if current is None:
                    current = {"header": "", "bullets": []}
                current["bullets"].append(line[2:].strip())
            else:
                if current:
                    sections.append(current)
                current = {"header": line, "bullets": []}
        if current:
            sections.append(current)
        return sections

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    base_style = doc.styles["Normal"]
    base_style.font.name = "Times New Roman"
    base_style.font.size = Pt(11)

    def set_run_font(run, *, bold=False, italic=False, size=11):
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic

    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        set_run_font(run, bold=True, size=11)
        border = p._element.get_or_add_pPr()
        from docx.oxml import OxmlElement
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "single")
        bottom.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz", "6")
        bottom.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space", "1")
        bottom.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color", "auto")
        p_bdr.append(bottom)
        border.append(p_bdr)

    def add_bullet_line(text, indent_inches=0.2):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(indent_inches)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=11)

    def add_left_right_line(left_text, right_text="", *, left_bold=False, left_italic=False, right_italic=False, bullet=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
        left_prefix = u"\u2022 " if bullet else ""
        left_run = p.add_run(f"{left_prefix}{left_text}")
        set_run_font(left_run, bold=left_bold, italic=left_italic, size=11)
        if right_text:
            right_run = p.add_run(f"\t{right_text}")
            set_run_font(right_run, italic=right_italic, size=11)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(profile.get("name", "Your Name"))
    set_run_font(name_run, bold=True, size=16)

    contact = " | ".join(filter(None, [
        profile.get("email"),
        profile.get("phone"),
        profile.get("linkedin"),
        profile.get("portfolio")
    ]))
    if contact:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(8)
        contact_run = contact_p.add_run(contact)
        set_run_font(contact_run, size=10)

    if education_school or education_grad or education_degree:
        add_section_heading("EDUCATION")
        add_left_right_line(education_school, f"Graduation {education_grad}" if education_grad else "", left_italic=True, right_italic=True)
        if education_degree:
            add_bullet_line(education_degree)

    skill_lines = [line.strip() for line in tailored_skills.splitlines() if line.strip()]
    if skill_lines:
        add_section_heading("TECHNICAL SKILLS")
        cleaned_skills = []
        for line in skill_lines:
            clean = line[2:].strip() if line.startswith("- ") else line
            category, values = (clean.split(":", 1) + [""])[:2] if ":" in clean else ("", clean)
            cleaned_skills.append({"category": category.strip(), "values": values.strip()})

        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        left_skills = cleaned_skills[::2]
        right_skills = cleaned_skills[1::2]
        row_total = max(len(left_skills), len(right_skills))

        for i in range(row_total):
            row = table.add_row().cells
            for cell in row:
                cell.width = Inches(3.6)

            for cell, entry in ((row[0], left_skills[i] if i < len(left_skills) else None), (row[1], right_skills[i] if i < len(right_skills) else None)):
                cell.text = ""
                if not entry:
                    continue
                p = cell.paragraphs[0]
                p.style = "List Bullet"
                p.paragraph_format.space_after = Pt(0)
                if entry["category"]:
                    run = p.add_run(f"{entry['category']}:")
                    set_run_font(run, bold=True, size=10)
                    value_run = p.add_run(f" {entry['values']}")
                    set_run_font(value_run, size=10)
                else:
                    run = p.add_run(entry["values"])
                    set_run_font(run, size=10)

    cert_entries = certifications_entries if certifications_entries else parse_certifications_text_to_entries(certifications_text)
    if cert_entries:
        add_section_heading("CERTIFICATIONS")
        for entry in cert_entries:
            name = normalize_text_block(entry.get("name"))
            date = normalize_text_block(entry.get("date"))
            description = normalize_text_block(entry.get("description"))
            if name or date:
                add_left_right_line(name, date, left_bold=True, right_italic=True, bullet=True)
            if description:
                desc = doc.add_paragraph()
                desc.paragraph_format.left_indent = Inches(0.2)
                desc.paragraph_format.space_after = Pt(2)
                run = desc.add_run(description)
                set_run_font(run, size=10.5)

    project_sections = parse_structured_text(tailored_projects)
    if project_sections:
        add_section_heading("PROJECTS")
        for entry in project_sections:
            header = entry["header"]
            bullets = entry["bullets"]
            if header:
                header_parts = [part.strip() for part in header.split("|")]
                title = header_parts[0] if header_parts else ""
                details = " | ".join(header_parts[1:]) if len(header_parts) > 1 else ""
                title_p = doc.add_paragraph()
                title_p.paragraph_format.space_after = Pt(0)
                title_run = title_p.add_run(title)
                set_run_font(title_run, bold=True, size=11)
                if details:
                    detail_p = doc.add_paragraph()
                    detail_p.paragraph_format.space_after = Pt(0)
                    detail_run = detail_p.add_run(details)
                    set_run_font(detail_run, italic=True, size=10)
            for bullet in bullets:
                add_bullet_line(bullet)

    experience_sections = parse_structured_text(tailored_experience)
    if experience_sections:
        add_section_heading("WORK EXPERIENCE")
        for entry in experience_sections:
            header = entry["header"]
            bullets = entry["bullets"]
            if header:
                header_parts = [part.strip() for part in header.split("|") if part.strip()]
                left_text = ""
                right_text = ""
                if header_parts:
                    if len(header_parts) == 1:
                        left_text = header_parts[0]
                    elif len(header_parts) == 2 and not looks_like_date_text(header_parts[1]):
                        left_text = ", ".join(header_parts)
                    else:
                        left_text = ", ".join(header_parts[:-1])
                        right_text = header_parts[-1]
                add_left_right_line(left_text, right_text, left_bold=True, right_italic=True)
            for bullet in bullets:
                add_bullet_line(bullet)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=build_export_filename(profile.get("name", ""), "docx"),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.route("/resumes")
def resumes():
    if "user_id" not in session:
        return redirect(url_for("login"))

    user_id = session["user_id"]
    selected_folder = request.args.get("folder", "All Resumes")

    all_resumes = list(
        resumes_collection.find({"user_id": user_id}).sort("updated_at", -1)
    )
    folder_items = build_folder_sidebar(user_id, all_resumes)
    folder_names = [item["name"] for item in folder_items]

    if selected_folder == "All Resumes":
        filtered_resumes = all_resumes
    else:
        filtered_resumes = [
            resume for resume in all_resumes
            if resume.get("folder", "Saved Drafts") == selected_folder
        ]

    for resume in filtered_resumes:
        display_date, display_time = format_datetime_for_display(resume.get("updated_at"))
        resume["display_date"] = display_date
        resume["display_time"] = display_time

    return render_template(
        "resumes.html",
        resumes=filtered_resumes,
        folder_items=folder_items,
        folder_names=folder_names,
        selected_folder=selected_folder
    )


@app.route("/create-folder", methods=["POST"])
def create_folder():
    if "user_id" not in session:
        return redirect(url_for("login"))

    user_id = session["user_id"]
    folder_name = request.form.get("folder_name", "").strip()

    if not folder_name:
        return redirect(url_for("resumes"))

    ensure_folder_record(user_id, folder_name)

    return redirect(url_for("resumes", folder=folder_name))


@app.route("/folder-action", methods=["POST"])
def folder_action():
    if "user_id" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401

    user_id = session["user_id"]
    folder_name = request.form.get("folder_name", "").strip()
    action = request.form.get("action", "").strip()
    new_name = request.form.get("new_name", "").strip()

    if not folder_name or folder_name == "Saved Drafts":
        return jsonify({"success": False, "error": "Folder cannot be modified."}), 400

    ensure_folder_record(user_id, folder_name)
    folder_docs = list(
        folders_collection.find({"user_id": user_id}).sort("sort_order", 1)
    )
    target_index = next((index for index, item in enumerate(folder_docs) if item.get("name") == folder_name), None)

    if target_index is None:
        return jsonify({"success": False, "error": "Folder not found."}), 404

    if action == "rename":
        if not new_name or new_name == "Saved Drafts":
            return jsonify({"success": False, "error": "Choose a different folder name."}), 400

        existing_folder = folders_collection.find_one({
            "user_id": user_id,
            "name": new_name
        })
        if existing_folder:
            return jsonify({"success": False, "error": "A folder with that name already exists."}), 400

        folders_collection.update_one(
            {"user_id": user_id, "name": folder_name},
            {"$set": {"name": new_name}}
        )
        resumes_collection.update_many(
            {"user_id": user_id, "folder": folder_name},
            {"$set": {"folder": new_name, "updated_at": datetime.utcnow()}}
        )
        return jsonify({"success": True, "folder": new_name})

    if action == "delete":
        resumes_collection.update_many(
            {"user_id": user_id, "folder": folder_name},
            {"$set": {"folder": "Saved Drafts", "updated_at": datetime.utcnow()}}
        )
        folders_collection.delete_one({"user_id": user_id, "name": folder_name})
        return jsonify({"success": True})

    if action not in {"move_up", "move_down"}:
        return jsonify({"success": False, "error": "Unsupported action."}), 400

    swap_index = target_index - 1 if action == "move_up" else target_index + 1
    if swap_index < 0 or swap_index >= len(folder_docs):
        return jsonify({"success": False, "error": "Folder cannot move further."}), 400

    current_folder = folder_docs[target_index]
    swap_folder = folder_docs[swap_index]
    current_order = current_folder.get("sort_order", target_index + 1)
    swap_order = swap_folder.get("sort_order", swap_index + 1)

    folders_collection.update_one(
        {"_id": current_folder["_id"]},
        {"$set": {"sort_order": swap_order}}
    )
    folders_collection.update_one(
        {"_id": swap_folder["_id"]},
        {"$set": {"sort_order": current_order}}
    )

    return jsonify({"success": True})


@app.route("/move-resume/<resume_id>", methods=["POST"])
def move_resume(resume_id):
    if "user_id" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401

    user_id = session["user_id"]
    target_folder = request.form.get("folder", "").strip() or "Saved Drafts"

    ensure_folder_record(user_id, target_folder)

    result = resumes_collection.update_one(
        {
            "_id": ObjectId(resume_id),
            "user_id": user_id
        },
        {
            "$set": {
                "folder": target_folder,
                "updated_at": datetime.utcnow()
            }
        }
    )

    if result.matched_count == 1:
        return jsonify({"success": True, "folder": target_folder})

    return jsonify({"success": False, "error": "Resume not found"}), 404


@app.route("/rename-resume/<resume_id>", methods=["POST"])
def rename_resume(resume_id):
    if "user_id" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401

    new_title = request.form.get("title", "").strip()
    if not new_title:
        return jsonify({"success": False, "error": "Resume title is required."}), 400

    user_id = session["user_id"]
    result = resumes_collection.update_one(
        {
            "_id": ObjectId(resume_id),
            "user_id": user_id
        },
        {
            "$set": {
                "title": new_title,
                "updated_at": datetime.utcnow()
            }
        }
    )

    if result.matched_count == 1:
        return jsonify({"success": True, "title": new_title})

    return jsonify({"success": False, "error": "Resume not found"}), 404


@app.route("/duplicate-resume/<resume_id>", methods=["POST"])
def duplicate_resume(resume_id):
    if "user_id" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401

    user_id = session["user_id"]
    existing_resume = resumes_collection.find_one({
        "_id": ObjectId(resume_id),
        "user_id": user_id
    })

    if not existing_resume:
        return jsonify({"success": False, "error": "Resume not found"}), 404

    duplicate_doc = {key: value for key, value in existing_resume.items() if key != "_id"}
    original_title = duplicate_doc.get("title") or "Untitled Resume"
    duplicate_doc["title"] = f"{original_title} (Copy)"
    duplicate_doc["created_at"] = datetime.utcnow()
    duplicate_doc["updated_at"] = datetime.utcnow()

    result = resumes_collection.insert_one(duplicate_doc)
    return jsonify({"success": True, "resume_id": str(result.inserted_id)})


@app.route("/delete-resume/<resume_id>", methods=["POST"])
def delete_resume(resume_id):
    if "user_id" not in session:
        return jsonify({"success": False}), 401

    user_id = session["user_id"]

    result = resumes_collection.delete_one({
        "_id": ObjectId(resume_id),
        "user_id": user_id
    })

    if result.deleted_count == 1:
        return jsonify({"success": True})

    return jsonify({"success": False}), 404

@app.route("/profile", methods=["GET"])
def profile():
    if "user_id" not in session:
        return redirect(url_for("login"))

    user_id = session["user_id"]
    profile_data = prepare_profile_for_view(profiles_collection.find_one({"user_id": user_id}))

    return render_template("profile.html", profile=profile_data)


@app.route("/save-profile", methods=["POST"])
def save_profile():
    if "user_id" not in session:
        return redirect(url_for("login"))

    user_id = session["user_id"]
    skills_entries = normalize_skills_entries(request.form.get("skills_entries"))
    projects_entries = normalize_resume_entries(request.form.get("projects_entries"))
    experience_entries = normalize_resume_entries(request.form.get("experience_entries"))
    certifications_entries = normalize_certification_entries(request.form.get("certifications_entries"))

    profile_doc = {
        "user_id": user_id,
        "name": request.form.get("fullname"),
        "email": request.form.get("email"),
        "phone": request.form.get("phone"),
        "location": request.form.get("location"),
        "linkedin": request.form.get("linkedin"),
        "github": request.form.get("github"),
        "portfolio": request.form.get("portfolio"),

        "school": request.form.get("school"),
        "school_location": request.form.get("school_location"),
        "expected_grad": request.form.get("expected_grad"),
        "degree": request.form.get("degree"),

        "skills": format_skills_entries(skills_entries),
        "skills_entries": skills_entries,
        "projects": format_resume_entries(projects_entries),
        "experience": format_resume_entries(experience_entries),
        "projects_entries": projects_entries,
        "experience_entries": experience_entries,
        "certifications": format_certification_entries(certifications_entries),
        "certifications_entries": certifications_entries,
        "updated_at": datetime.utcnow()
    }

    profiles_collection.update_one(
        {"user_id": user_id},
        {"$set": profile_doc},
        upsert=True
    )

    return redirect(url_for("profile"))


@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        name = request.form.get("name")
        email = request.form.get("email")
        password = request.form.get("password")

        existing_user = users_collection.find_one({"email": email})
        if existing_user:
            return "An account with that email already exists."

        hashed_password = generate_password_hash(password)

        new_user = {
            "name": name,
            "email": email,
            "password": hashed_password,
            "created_at": datetime.utcnow()
        }

        result = users_collection.insert_one(new_user)

        session["user_id"] = str(result.inserted_id)
        session["user_name"] = name

        return redirect(url_for("dashboard"))

    return render_template("signup.html")


@app.route("/privacy-policy")
def privacy_policy():
    return render_template("privacy_policy.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email")
        password = request.form.get("password")

        user = users_collection.find_one({"email": email})

        if user and check_password_hash(user["password"], password):
            session["user_id"] = str(user["_id"])
            session["user_name"] = user["name"]
            return redirect(url_for("dashboard"))

        return "Invalid email or password."

    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/job-tracker")
def job_tracker():
    if "user_id" not in session:
        return redirect(url_for("login"))

    user_id = session["user_id"]
    jobs = list(jobs_collection.find({"user_id": user_id}).sort("updated_at", -1))

    grouped_jobs = {
        "Saved": [],
        "Applied": [],
        "Interview": [],
        "Offer": [],
        "Rejected": []
    }

    for job in jobs:
        status = job.get("status", "Saved")
        created_date, created_time = format_datetime_for_display(job.get("created_at"))
        updated_date, updated_time = format_datetime_for_display(job.get("updated_at"))
        job["display_created_date"] = created_date
        job["display_created_time"] = created_time
        job["display_updated_date"] = updated_date
        job["display_updated_time"] = updated_time
        if status in grouped_jobs:
            grouped_jobs[status].append(job)

    return render_template("job_tracker.html", grouped_jobs=grouped_jobs)


@app.route("/add-job", methods=["POST"])
def add_job():
    if "user_id" not in session:
        return redirect(url_for("login"))

    user_id = session["user_id"]

    new_job = {
        "user_id": user_id,
        "company": request.form.get("company"),
        "job_title": request.form.get("job_title"),
        "location": request.form.get("location"),
        "status": request.form.get("status"),
        "resume_name": request.form.get("resume_name"),
        "job_link": request.form.get("job_link"),
        "notes": request.form.get("notes"),
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow()
    }

    jobs_collection.insert_one(new_job)
    return redirect(url_for("job_tracker"))


@app.route("/update-job-status/<job_id>", methods=["POST"])
def update_job_status(job_id):
    if "user_id" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401

    user_id = session["user_id"]
    new_status = (request.form.get("status") or "").strip()
    valid_statuses = {"Saved", "Applied", "Interview", "Offer", "Rejected"}

    if new_status not in valid_statuses:
        return jsonify({"success": False, "error": "Invalid status."}), 400

    result = jobs_collection.update_one(
        {
            "_id": ObjectId(job_id),
            "user_id": user_id
        },
        {
            "$set": {
                "status": new_status,
                "updated_at": datetime.utcnow()
            }
        }
    )

    if result.matched_count == 1:
        return jsonify({"success": True, "status": new_status})

    return jsonify({"success": False, "error": "Job not found."}), 404


@app.route("/delete-job/<job_id>", methods=["POST"])
def delete_job(job_id):
    if "user_id" not in session:
        return jsonify({"success": False}), 401

    user_id = session["user_id"]

    result = jobs_collection.delete_one({
        "_id": ObjectId(job_id),
        "user_id": user_id
    })

    if result.deleted_count == 1:
        return jsonify({"success": True})

    return jsonify({"success": False}), 404

if __name__ == "__main__":
    app.run(host=HOST, port=PORT)
